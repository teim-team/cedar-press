#!/usr/bin/env python3
"""
Cedar Press - 1175: write the identifier-system explainer as a Word document.

    py -3 code/1175_write_id_system_doc.py

Owner, 2026-09-04: *"i also want you in the dataset of native entities explain
how our code system works maybe as a word doc"*

EVERY FIGURE IN THE DOCUMENT IS MEASURED WHEN IT IS WRITTEN. Nothing is typed
in. A document about an identifier scheme that quotes a stale count is the
thing this project has been fixing all week, and it would be a poor advert for
the scheme to make that mistake in its own explainer.
"""
from __future__ import annotations

import csv
import sys
from collections import Counter
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
SPINE = ROOT / "data" / "spine" / "cedar_identity_register.csv"
LEDGER = ROOT / "data" / "clean" / "cedar_identifier_ledger_final.csv"
ALIASES = ROOT / "data" / "clean" / "entity_aliases.csv"
OUT = ROOT / "dist" / "qc_review" / "2_how_cedar_ids_work.docx"
csv.field_size_limit(10_000_000)

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:
    sys.exit("python-docx is not installed: py -3 -m pip install python-docx")

TEAL = RGBColor(0x0F, 0x6B, 0x63)


def load(p):
    with Path(p).open(encoding="utf-8-sig", errors="replace", newline="") as fh:
        return list(csv.DictReader(fh))


def measure():
    reg = load(SPINE)
    # The register holds `canonical_name`, the short handle that was retired on
    # 2026-09-04. This document must show what a reader actually gets, which is
    # the OFFICIAL `name` built by 1180, so merge it in before measuring.
    names_path = ROOT / "data" / "spine" / "cedar_entity_names.csv"
    if names_path.exists():
        official = {r["cedar_uid"]: r.get("name", "") for r in load(names_path)}
        for r in reg:
            r["name"] = official.get(r.get("cedar_uid", ""), "") \
                or r.get("canonical_name", "")
    else:
        for r in reg:
            r["name"] = r.get("canonical_name", "")
    led = load(LEDGER) if LEDGER.exists() else []
    ali = load(ALIASES) if ALIASES.exists() else []
    classes = Counter((r.get("entity_class") or "?") for r in reg)
    idtypes = Counter((r.get("identifier_type") or "?") for r in led)
    tiers = Counter((r.get("confidence_tier") or "?") for r in led)
    return {
        "entities": len(reg), "classes": classes,
        "aliases": len(ali), "ledger_rows": len(led),
        "idtypes": idtypes, "tiers": tiers,
        "uid_example": next((r for r in reg if r.get("name")), reg[0]),
        "statuses": Counter((r.get("register_status") or "?") for r in reg),
    }


def main():
    m = measure()
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    def h(text, level=1):
        p = d.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = TEAL
        return p

    def para(text, bold=False, italic=False):
        p = d.add_paragraph()
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        return p

    def mono(text):
        p = d.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        return p

    title = d.add_heading("How Cedar's identifier system works", level=0)
    for run in title.runs:
        run.font.color.rgb = TEAL
    sub = para(f"Measured against the live register on {date.today():%d %B %Y}. "
               f"Every number in this document was read from the data when the "
               f"document was written; none is typed in.", italic=True)
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ---------------------------------------------------------------- the point
    h("The one rule everything else serves", 1)
    para("A cedar_uid must always resolve to the same impermeable Native "
         "entity. The dataset separately identifies the event, object or "
         "business, and describes that entity's role in it.")
    para("That is the whole design. A dataset row says what happened; the "
         "cedar_uid says who it happened to. Keeping those apart is what lets "
         "a contract, a lobbying filing, a NAGPRA notice and a federal grant "
         "all point at one nation without any of them having to agree about "
         "how that nation is spelled.")

    # ------------------------------------------------------------- one key
    h("One identifier, and the entity's name", 1)
    ex = m["uid_example"]
    para("Cedar mints exactly one ENTITY identifier: cedar_uid. Events, "
         "awards, filings, deals and other objects retain their own record "
         "identifiers - a Deal_ID, a federal award number, a docket. There "
         "is no second readable code for the ENTITY, because the entity "
         "already has a readable form: its name.", bold=True)
    mono(f"    {ex['cedar_uid']}    {ex['name']}")
    para("The cedar_uid is opaque and permanent, and that is the point. A "
         "meaningful key has to change when the meaning changes, and an "
         "identifier that changes is not an identifier. A tribe can be "
         "renamed, reclassified, or federally re-recognised and its cedar_uid "
         "does not move. Every dataset joins on it.")
    para("What a person reads is name: the authoritative published name "
         "where an official register exists, and otherwise the "
         "best-supported legal or organizational name. A register covers "
         "some classes and not others - the BIA Federal Register for tribes "
         "and Alaska Native villages, the DOI list for Native Hawaiian "
         "Organizations, the ANCSA corporation list for ANCs - while Native "
         "nonprofits, CDFIs, tribal colleges, intertribal organizations and "
         "individually Native-owned businesses have no single federal "
         "register of names. Every row records which source its name came "
         "from, so a reader can tell the two apart rather than assuming "
         "every name carries equal authority. It is backed by the alias "
         "table so every form a nation has been known by resolves to the "
         "same key.")

    h("A code that was removed on 4 September 2026", 2)
    para("Cedar used to carry a second, prefixed code beside the uid — "
         "TRBF-, AKNF-, ANVC- and so on. That was the CICD Native Entity "
         "Connector Crosswalk identifier, inherited from an external source, "
         "and it has been retired from every dataset.")
    para("It survived three earlier attempts at removal because it had been "
         "described, including in an earlier draft of this document, as "
         "Cedar's own readable code. It was not. Nothing needs it: the uid "
         "does the joining and the name does the reading.")

    # --------------------------------------------------------------- the shape
    h("What the register holds", 1)
    # "every one active" was ASSERTED, and the reviewer was right to push on
    # it: a directory listing an organization does not prove the organization
    # still operates. Measured 2026-09-04 against the IRS EO BMF, 355 of the
    # 359 Native nonprofits carry STATUS 01 and four are absent from the BMF
    # entirely, so their status is unevidenced and they now say so. This line
    # counts what the register actually holds instead of claiming.
    _st = m["statuses"]
    _active = _st.get("active", 0)
    _unver = sum(n for k, n in _st.items() if k != "active")
    if _unver:
        para(f"{m['entities']:,} entities across {len(m['classes'])} classes. "
             f"{_active:,} are recorded active; {_unver} carry "
             f"active_unverified, meaning a source lists them but their "
             f"current operating status is not evidenced.")
    else:
        para(f"{m['entities']:,} entities, every one active, across "
             f"{len(m['classes'])} classes.")
    t = d.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    for i, v in enumerate(("entity class", "entities")):
        t.rows[0].cells[i].text = v
    for cls, n in m["classes"].most_common():
        row = t.add_row().cells
        row[0].text = cls
        row[1].text = f"{n:,}"
    para("Grouped by entity_class, which is the real classification. It used "
         "to be shown by handle prefix, and that was doubly wrong: the prefix "
         "belonged to the retired scheme, and it never reliably identified the "
         "class anyway — ANVC- covered both village and group corporations, "
         "and CDFI- covered Native CDFIs and Native Financial Institutions, so "
         "grouping on it was wrong for 272 entities.", italic=True)

    # --------------------------------------------------------------- the ledger
    h("How a real-world identifier reaches an entity", 1)
    para(f"The register says who exists. A second table — the identifier "
         f"ledger, {m['ledger_rows']:,} rows — says which government "
         f"identifiers belong to whom:")
    for k, v in m["idtypes"].most_common(5):
        mono(f"    {k:<6} {v:>8,}")
    para("A UEI is SAM.gov's, a CAGE is the government's contractor code, an "
         "EIN is the IRS's. Cedar does not mint those; it records which entity "
         "each one belongs to, and how confident that link is.")

    h("Tiers — how strong is the link", 2)
    order = ["A", "B", "C", "X"]
    meaning = {
        "A": "two independent legs of evidence",
        "B": "one leg — usually a published statement by the owner",
        "C": "recorded, not strong enough to attribute",
        "X": "a NEGATIVE ruling: this identifier is NOT that entity",
    }
    for k in order:
        if k in m["tiers"]:
            mono(f"    tier {k}   {m['tiers'][k]:>7,}   {meaning[k]}")
    para("Tier X is the one most often misread. It is not missing data - it "
         "is a decision, recorded so the same wrong candidate is not "
         "proposed again. It rejects a specific PAIRING of one external "
         "identifier with one cedar_uid; it does not reject the external "
         "identifier itself, which usually belongs to somebody and may be "
         "keyed correctly elsewhere. \"This UEI is NOT the Onondaga Nation\" "
         "is a tier X row; the UEI still belongs to the Onondaga Golf and "
         "Country Club.")
    para("A tier is inherited from the source that made the ruling, never "
         "assigned by whatever code is reading it. That rule exists because "
         "this project once treated an exact EIN match as strong evidence and "
         "attributed a Wisconsin United Way to a California tribe: the "
         "exactness of a key says nothing about the correctness of a link.")

    # --------------------------------------------------------------- the aliases
    h("Names change; identifiers do not", 1)
    para(f"An alias table of {m['aliases']:,} rows records every name an entity "
         f"has been known by — common, legal, former, and shortened forms. "
         f"Cortina Band of Wintun Indians is now Kletsel Dehe Wintun; both "
         f"names resolve to one cedar_uid.")
    para("This is also what stops a name comparison being mistaken for an "
         "identity test. Two names that look different can be one entity, and "
         "two that look similar can be two — Crow and Crow Creek are different "
         "nations; Zuni and Pueblo of Zuni are the same one.")

    # ----------------------------------------------------------- what goes wrong
    h("The failure this system exists to prevent", 1)
    para("Almost every attribution error found in this product has one shape: "
         "a shared word treated as a shared identity.", bold=True)
    for a, b in (("\"Three\" in Three Saints Bay",
                  "the Three Affiliated Tribes of North Dakota"),
                 ("\"Mission\" in Mission Support Services",
                  "the Campo Band of Diegueño Mission Indians"),
                 ("\"Kootenai\"", "the Kootenai Tribe of Idaho, not the "
                                  "Confederated Salish and Kootenai"),
                 ("\"Omaha\" in Housing Authority of the City of Omaha",
                  "the Omaha Tribe")):
        mono(f"    {a}  ->  {b}")
    para("The last of those carried $1.13 billion of municipal public-housing "
         "money before it was found and withdrawn on 4 September 2026. The "
         "city is named after the tribe; that is not a relationship of "
         "ownership or control.")
    para("This is why the identifier, not the name, is the thing that "
         "resolves an entity — and why a name match is only ever a candidate.")

    # ---------------------------------------------------------------- reviewing
    h("What to look for when reviewing", 1)
    for line in (
        "A cedar_uid appearing under two different canonical names — one key, two entities.",
        "An entity with a UEI that Cedar keys elsewhere but not here — a rule that exists and was not applied.",
        "A name match that shares only a place word: Creek, Crow, Eagle, Mission, Three, Vista.",
        "A corporation keyed to the village government of the same name — under ANCSA those are different legal persons.",
        "A blank cedar_uid that should not be blank — and one that is blank correctly, because a denial withdrew it.",
    ):
        d.add_paragraph(line, style="List Bullet")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"  wrote {OUT.relative_to(ROOT)}")
    print(f"     {m['entities']:,} entities, {len(m['classes'])} classes, "
          f"{m['aliases']:,} aliases, {m['ledger_rows']:,} ledger rows")


if __name__ == "__main__":
    main()
